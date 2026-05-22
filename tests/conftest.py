import zipfile
from pathlib import Path
from typing import Any, Dict
import fitz
import pytest
from docx import Document
from PIL import Image, ImageDraw

@pytest.fixture
def sample_images(tmp_path: Path) -> Dict[str, Path]:
    assets = tmp_path / "assets"
    assets.mkdir()

    big_png = assets / "big_simple.png"
    small_icon = assets / "small_icon.png"
    photo_jpg = assets / "photo.jpg"

    # Imagen grande y comprimible
    img = Image.new("RGB", (900, 600), (245, 245, 245))
    draw = ImageDraw.Draw(img)
    draw.rectangle([40, 40, 860, 560], outline=(0, 0, 0), width=6)
    draw.text((60, 60), "BIG PNG", fill=(0, 0, 0))
    img.save(big_png, format="PNG", optimize=True)

    # Icono Pequeño
    icon = Image.new("RGB", (32, 32), (255, 255, 255))
    draw = ImageDraw.Draw(icon)
    draw.rectangle([2, 2, 29, 29], outline=(0, 0, 0), width=2)
    draw.line([6, 16, 26, 16], fill=(0, 0, 0), width=2)
    draw.line([16, 6, 16, 26], fill=(0, 0, 0), width=2)
    icon.save(small_icon, format="PNG", compress_level=0)

    # Imagen JPG Mediana
    jpg = Image.new("RGB", (700, 450), (255, 255, 255))
    draw = ImageDraw.Draw(jpg)
    for i in range(20):
        x0 = 20 + i * 12
        y0 = 40 + i * 8
        x1 = x0 + 160
        y1 = y0 + 80
        draw.rectangle([x0, y0, x1, y1], outline=(30, 90, 180), width=3)
    draw.text((40, 40), "PHOTO JPG", fill=(0, 0, 0))
    jpg.save(photo_jpg, format="JPEG", quality=85, optimize=True)

    return {
        "big_png": big_png,
        "small_icon": small_icon,
        "photo_jpg": photo_jpg,
    }

@pytest.fixture
def base_cfg(tmp_path: Path) -> Dict[str, Dict]:
    return {
        "paths": {
            "input_dir": str(tmp_path / "input"),
            "output_dir": str(tmp_path / "output"),
            "temp_dir": str(tmp_path / "temp"),
        },
        "filters": {
            "min_kb": 1,
            "min_width": 0,
            "min_height": 0,
        },
        "dedup": {
            "enabled": True,
        },
        "output": {
            "format": "zip",
        },
        "logging": {
            "level": "INFO",
            "log_file": str(tmp_path / "logs" / "run.log"),
            "format": "%(asctime)s | %(levelname)s | %(message)s",
            "datefmt": "%Y-%m-%d %H:%M:%S",
        },
    }

# Documento de prueba DOCX
def create_docx(path: Path, image_paths: list[Path]) -> Path:
    doc = Document()
    doc.add_heading("TEST DOCX", level=1)

    for image_path in image_paths:
        doc.add_paragraph(image_path.name)
        doc.add_picture(str(image_path))
    doc.save(str(path))
    return path

# Documento de prueba PDF
def create_pdf(path: Path, image_path: Path) -> Path:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    rect = fitz.Rect(72, 72, 420, 360)
    page.insert_image(rect, stream=image_path.read_bytes())
    doc.save(str(path))
    doc.close()
    return path


# Documento de prueba PPTX
def create_fake_pptx(path: Path, media: dict[str, bytes]) -> Path:
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", "<Types></Types>")

        for filename, data in media.items():
            zf.writestr(f"ppt/media/{filename}", data)

    return path


# Documento de prueba XLSX
def create_fake_xlsx(path: Path, media: dict[str, bytes]) -> Path:
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", "<Types></Types>")

        for filename, data in media.items():
            zf.writestr(f"xl/media/{filename}", data)

    return path


# Creacion XML Base
def create_empty_ooxml(path: Path) -> Path:
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", "<Types></Types>")

    return path


def zip_names(zip_path: Path) -> list[str]:
    with zipfile.ZipFile(zip_path, "r") as zf:
        return zf.namelist()